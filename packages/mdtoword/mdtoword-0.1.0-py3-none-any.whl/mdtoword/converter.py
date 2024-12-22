from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc import constants as opc_constants
import markdown
from bs4 import BeautifulSoup
import os

class MarkdownToDocxConverter:
    """A class to convert Markdown files to Word documents with advanced formatting."""

    # Estilos por defecto
    DEFAULT_STYLES = {
        'document': {
            'font_name': 'Calibri',
            'font_size': 11,
            'margins': {
                'top': 1.0,
                'bottom': 1.0,
                'left': 1.0,
                'right': 1.0
            }
        },
        'h1': {
            'font_name': 'Arial',
            'font_size': 16,
            'color': (0, 0, 0),
            'bold': True
        },
        'h2': {
            'font_name': 'Arial',
            'font_size': 14,
            'color': (0, 0, 0),
            'bold': True
        },
        'h3': {
            'font_name': 'Arial',
            'font_size': 12,
            'color': (0, 0, 0),
            'bold': True
        },
        'paragraph': {
            'font_name': 'Calibri',
            'font_size': 11,
            'color': (0, 0, 0)
        },
        'code': {
            'font_name': 'Courier New',
            'font_size': 10,
            'color': (169, 169, 169)
        },
        'strong': {
            'bold': True,
            'color': (0, 0, 0)
        },
        'emphasis': {
            'italic': True,
            'color': (0, 0, 0)
        }
    }

    def __init__(self, md_file=None, docx_file=None, template_file=None, styles=None, keep_template_styles=True):
        """
        Initialize the converter.
        
        Args:
            md_file (str, optional): Path to the input Markdown file
            docx_file (str, optional): Path to the output Word document
            template_file (str, optional): Path to a template Word document
            styles (dict, optional): Custom styles for the document
            keep_template_styles (bool): Whether to keep template styles when using a template
        """
        # Inicializar estilos
        self.styles = self.DEFAULT_STYLES.copy()
        if styles:
            for key, value in styles.items():
                if key in self.styles:
                    self.styles[key].update(value)
        
        # Crear documento desde plantilla o nuevo
        if template_file:
            self.doc = Document(template_file)
            if not keep_template_styles:
                self._clear_content_keeping_headers_footers()
        else:
            self.doc = Document()
            self._setup_default_document()

        self.md_file = md_file
        self.docx_file = docx_file

    def convert_file(self, input_file=None, output_file=None):
        """
        Convert a Markdown file to a Word document.
        
        Args:
            input_file (str, optional): Path to the input Markdown file
            output_file (str, optional): Path where the output Word document will be saved
        """
        md_file = input_file or self.md_file
        docx_file = output_file or self.docx_file
        
        if not md_file or not docx_file:
            raise ValueError("Input and output files are required")

        try:
            with open(md_file, 'r', encoding='utf-8') as f:
                markdown_text = f.read()
            self.convert_string(markdown_text, docx_file)
            return True
        except Exception as e:
            print(f"Error during conversion: {str(e)}")
            return False

    def convert_string(self, markdown_text, output_file):
        """
        Convert a Markdown string to a Word document.
        
        Args:
            markdown_text (str): Markdown formatted text
            output_file (str): Path where the output Word document will be saved
        """
        # Convert markdown to HTML
        html = markdown.markdown(
            markdown_text,
            extensions=[
                'fenced_code',
                'tables',
                'nl2br',
                'sane_lists'
            ]
        )
        soup = BeautifulSoup(html, 'html.parser')
        
        # Process HTML elements
        for element in soup.children:
            self._process_element(element)
        
        # Remove extra empty paragraphs at the end
        while len(self.doc.paragraphs) > 0 and not self.doc.paragraphs[-1].text.strip():
            p = self.doc.paragraphs[-1]._element
            p.getparent().remove(p)
        
        # Save the document
        self.doc.save(output_file)

    def _setup_default_document(self):
        """Configure default document margins and sections."""
        section = self.doc.sections[0]
        margins = self.styles['document']['margins']
        section.top_margin = Inches(margins['top'])
        section.bottom_margin = Inches(margins['bottom'])
        section.left_margin = Inches(margins['left'])
        section.right_margin = Inches(margins['right'])

    def _clear_content_keeping_headers_footers(self):
        """Clear document content while keeping headers and footers."""
        # Save sections with their headers and footers
        sections_data = []
        for section in self.doc.sections:
            section_info = {
                'header': section.header,
                'footer': section.footer,
                'orientation': section.orientation,
                'page_height': section.page_height,
                'page_width': section.page_width,
                'margins': {
                    'top': section.top_margin,
                    'bottom': section.bottom_margin,
                    'left': section.left_margin,
                    'right': section.right_margin
                }
            }
            sections_data.append(section_info)

        # Clear main content
        self.doc._body.clear_content()

        # Restore sections with their properties
        for i, section_info in enumerate(sections_data):
            if i == 0:
                section = self.doc.sections[0]
            else:
                section = self.doc.add_section()
            
            section.orientation = section_info['orientation']
            section.page_height = section_info['page_height']
            section.page_width = section_info['page_width']
            section.top_margin = section_info['margins']['top']
            section.bottom_margin = section_info['margins']['bottom']
            section.left_margin = section_info['margins']['left']
            section.right_margin = section_info['margins']['right']

    def _apply_style(self, run, style_dict):
        """Apply style to a run."""
        try:
            if 'font_name' in style_dict:
                run.font.name = style_dict['font_name']
            
            if 'font_size' in style_dict:
                run.font.size = Pt(style_dict['font_size'])
            
            if 'color' in style_dict:
                run.font.color.rgb = RGBColor(*style_dict['color'])
            
            if 'bold' in style_dict:
                run.bold = style_dict['bold']
            
            if 'italic' in style_dict:
                run.italic = style_dict['italic']
                
        except Exception as e:
            print(f"Error applying style: {str(e)}")

    def _add_heading(self, text, level):
        """Add a heading with the corresponding style."""
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(text)
        style_key = f'h{level}'
        
        if style_key in self.styles:
            self._apply_style(run, self.styles[style_key])
        
        return paragraph

    def _process_element(self, element):
        """Process an HTML element and convert it to Word format."""
        if isinstance(element, str):
            if element.strip():
                p = self.doc.add_paragraph()
                run = p.add_run(element)
                self._apply_style(run, self.styles['paragraph'])
            return

        if element.name is None:
            return

        if element.name.startswith('h') and len(element.name) == 2:
            level = int(element.name[1])
            self._add_heading(element.text.strip(), level)
            self.doc.add_paragraph()
        
        elif element.name == 'p':
            if str(element).strip():
                p = self.doc.add_paragraph()
                self._process_inline_elements(element, p)
            self.doc.add_paragraph()
        
        elif element.name == 'hr':
            self._add_horizontal_line()
            self.doc.add_paragraph()
        
        elif element.name == 'table':
            self._add_table(element)
            self.doc.add_paragraph()
        
        elif element.name in ['ul', 'ol']:
            self._add_list(element, ordered=element.name == 'ol')
            self.doc.add_paragraph()
        
        elif element.name == 'pre':
            code = element.find('code')
            if code:
                p = self.doc.add_paragraph()
                run = p.add_run(code.get_text())
                self._apply_style(run, self.styles['code'])
            self.doc.add_paragraph()

    def _process_inline_elements(self, element, paragraph):
        """Process inline elements like bold, italic, etc."""
        for child in element.children:
            if isinstance(child, str):
                run = paragraph.add_run(child)
                self._apply_style(run, self.styles['paragraph'])
            else:
                if child.name == 'br':
                    paragraph.add_run('\n')
                    continue
                
                text = child.get_text()
                run = paragraph.add_run(text)
                
                if child.name in ['strong', 'b']:
                    self._apply_style(run, self.styles['strong'])
                elif child.name in ['em', 'i']:
                    self._apply_style(run, self.styles['emphasis'])
                elif child.name == 'code':
                    self._apply_style(run, self.styles['code'])
                elif child.name == 'a':
                    run.underline = True
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    # TODO: implement hyperlink

    def _add_table(self, table_element):
        """Add a table to the document."""
        rows = table_element.find_all('tr')
        if not rows:
            return
        
        num_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
        table = self.doc.add_table(rows=0, cols=num_cols)
        table.style = 'Table Grid'
        table.allow_autofit = True
        
        # Process header
        header_cells = rows[0].find_all(['th', 'td'])
        header_row = table.add_row()
        for j, cell in enumerate(header_cells):
            table_cell = header_row.cells[j]
            paragraph = table_cell.paragraphs[0]
            run = paragraph.add_run(cell.get_text().strip())
            run.bold = True
            self._apply_style(run, self.styles.get('table_header', self.styles['strong']))
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Process rows
        for row in rows[1:]:
            cells = row.find_all(['td', 'th'])
            table_row = table.add_row()
            for j, cell in enumerate(cells):
                table_cell = table_row.cells[j]
                paragraph = table_cell.paragraphs[0]
                run = paragraph.add_run(cell.get_text().strip())
                self._apply_style(run, self.styles.get('table_cell', self.styles['paragraph']))
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def _add_horizontal_line(self):
        """Add a horizontal line."""
        paragraph = self.doc.add_paragraph()
        p = paragraph._p
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        pBdr.append(bottom)
        p.get_or_add_pPr().append(pBdr)

    def _add_list(self, list_element, ordered=False):
        """Add an ordered or unordered list."""
        for item in list_element.find_all('li', recursive=False):
            paragraph = self.doc.add_paragraph()
            level = len(item.find_parents(['ul', 'ol'])) - 1
            
            if ordered:
                paragraph.style = 'List Number'
                paragraph._p.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = 1
            else:
                paragraph.style = 'List Bullet'
                paragraph._p.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = 2
            
            paragraph._p.get_or_add_pPr().get_or_add_numPr().get_or_add_ilvl().val = level
            self._process_inline_elements(item, paragraph)
